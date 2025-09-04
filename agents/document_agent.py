"""DocumentAgent - 文档处理Agent

负责解析和处理各种格式的文档。
"""

import os
from typing import List, Dict, Any, Optional
import logging

from .base import BaseAgent, Task, TaskResult, TaskStatus, AgentCapability

logger = logging.getLogger(__name__)


class DocumentAgent(BaseAgent):
    """文档处理Agent"""
    
    def __init__(self):
        super().__init__("DocumentAgent")
        self.supported_formats = [".pdf", ".docx", ".txt", ".md"]
        
    def get_capabilities(self) -> List[AgentCapability]:
        """获取DocumentAgent能力"""
        return [AgentCapability.DOCUMENT_PARSING]
    
    async def process(self, task: Task) -> TaskResult:
        """处理文档任务"""
        try:
            if task.task_type != "document_parsing":
                return TaskResult(
                    task_id=task.task_id,
                    status=TaskStatus.FAILED,
                    error=f"不支持的任务类型: {task.task_type}"
                )
            
            file_path = task.data.get("file_path")
            if not file_path:
                return TaskResult(
                    task_id=task.task_id,
                    status=TaskStatus.FAILED,
                    error="缺少文件路径"
                )
            
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return TaskResult(
                    task_id=task.task_id,
                    status=TaskStatus.FAILED,
                    error=f"文件不存在: {file_path}"
                )
            
            # 检查文件格式
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in self.supported_formats:
                return TaskResult(
                    task_id=task.task_id,
                    status=TaskStatus.FAILED,
                    error=f"不支持的文件格式: {file_ext}"
                )
            
            # 解析文档
            parsed_content = await self._parse_document(file_path, file_ext)
            
            # 保存解析结果到内存
            self.save_memory(f"document_{task.task_id}", parsed_content)
            
            logger.info(f"成功解析文档: {file_path}")
            return TaskResult(
                task_id=task.task_id,
                status=TaskStatus.COMPLETED,
                result=parsed_content
            )
            
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            return TaskResult(
                task_id=task.task_id,
                status=TaskStatus.FAILED,
                error=str(e)
            )
    
    async def _parse_document(self, file_path: str, file_ext: str) -> Dict[str, Any]:
        """解析文档内容"""
        if file_ext == ".pdf":
            return await self._parse_pdf(file_path)
        elif file_ext == ".docx":
            return await self._parse_docx(file_path)
        elif file_ext == ".txt":
            return await self._parse_txt(file_path)
        elif file_ext == ".md":
            return await self._parse_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}")
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文档"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    text_content += f"\n--- 第{page_num + 1}页 ---\n"
                    text_content += page.extract_text()
                
                return {
                    "file_path": file_path,
                    "file_type": "pdf",
                    "pages": len(pdf_reader.pages),
                    "text_content": text_content.strip(),
                    "metadata": {
                        "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                        "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                        "subject": pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else ''
                    }
                }
        except ImportError:
            raise ImportError("需要安装PyPDF2: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"PDF解析失败: {e}")
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取文本内容
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # 提取表格内容
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            return {
                "file_path": file_path,
                "file_type": "docx",
                "text_content": text_content.strip(),
                "tables": tables_content,
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables)
            }
        except ImportError:
            raise ImportError("需要安装python-docx: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word文档解析失败: {e}")
    
    async def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        """解析文本文档"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            text_content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                        used_encoding = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise Exception("无法解码文本文件")
            
            lines = text_content.split('\n')
            
            return {
                "file_path": file_path,
                "file_type": "txt",
                "text_content": text_content,
                "lines_count": len(lines),
                "characters_count": len(text_content),
                "encoding": used_encoding
            }
        except Exception as e:
            raise Exception(f"文本文件解析失败: {e}")
    
    async def _parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """解析Markdown文档"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # 简单的Markdown解析
            lines = content.split('\n')
            headers = []
            
            for line in lines:
                line = line.strip()
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    title = line.lstrip('#').strip()
                    headers.append({
                        "level": level,
                        "title": title
                    })
            
            return {
                "file_path": file_path,
                "file_type": "markdown",
                "text_content": content,
                "headers": headers,
                "lines_count": len(lines)
            }
        except Exception as e:
            raise Exception(f"Markdown文件解析失败: {e}")
    
    def get_supported_formats(self) -> List[str]:
        """获取支持的文件格式"""
        return self.supported_formats.copy()
    
    def is_supported_file(self, file_path: str) -> bool:
        """检查文件是否支持"""
        file_ext = os.path.splitext(file_path)[1].lower()
        return file_ext in self.supported_formats